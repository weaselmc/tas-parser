from docx import Document
from io import BytesIO
import re

from WordHtmlConverter import WordHtmlConverter


class LAPDoc:

    def __init__(self, source):

        self.source = source

        if isinstance(source, bytes):
            self.source = BytesIO(source)

        self.doc = Document(self.source)

        self.tables = self._load_tables()

        self.converter = WordHtmlConverter()

        self.lap = None
        self.sessions = []
        self.assessments = []
        self.lecturers = []
        self.session_element_mappings = []

    # =====================================================
    # CORE
    # =====================================================

    @staticmethod
    def clean(text):

        if text is None:
            return ""

        return " ".join(str(text).split())

    def _load_tables(self):

        tables = []

        for table_number, table in enumerate(self.doc.tables):

            rows = []

            for row in table.rows:

                rows.append([
                    self.clean(cell.text)
                    for cell in row.cells
                ])

            tables.append({
                "table_number": table_number,
                "table": table,
                "rows": rows
            })

        return tables

    def find_table_containing(self, search_text):

        search_text = search_text.lower()

        for table in self.tables:

            table_text = "\n".join(
                " ".join(row)
                for row in table["rows"]
            ).lower()

            if search_text in table_text:
                return table

        return None

    # =====================================================
    # QUALIFICATION
    # =====================================================

    def get_qualification_code(self):

        full_text = "\n".join(
            " ".join(row)
            for table in self.tables
            for row in table["rows"]
        )

        match = re.search(
            r"(ICT\d{5})",
            full_text
        )

        if match:
            return match.group(1)

        return None

    # =====================================================
    # CLUSTER
    # =====================================================

    def get_cluster_name(self):

        for table in self.tables:

            rows = table["rows"]

            for i, row in enumerate(rows):

                row_text = " ".join(row).lower()

                if "cluster name" in row_text:

                    for j in range(i + 1, len(rows)):

                        value = " ".join(
                            rows[j]
                        ).strip()

                        if value:
                            return value

        return None

    def make_cluster_key(self, cluster_name):

        if not cluster_name:
            return None

        cluster_name = cluster_name.lower()

        cluster_name = cluster_name.replace(
            "&",
            "and"
        )

        cluster_name = re.sub(
            r"[^a-z0-9]+",
            "-",
            cluster_name
        )

        cluster_name = re.sub(
            r"-+",
            "-",
            cluster_name
        )

        return cluster_name.strip("-")

    # =====================================================
    # LEARNING RESOURCES
    # =====================================================

    def get_learning_resources(self):

        table = self.find_table_containing(
            "student learning resources"
        )

        if table is None:
            return None

        html = self.converter.word_xml_to_html(
            table["table"]._tbl
        )

        #
        # Remove title row
        #
        html = re.sub(
            r'<tr>\s*<td><p>Student Learning Resources.*?</tr>',
            '',
            html,
            flags=re.IGNORECASE | re.DOTALL
        )

        #
        # Remove lecturer section and everything after it
        #
        html = re.sub(
            r'<tr>\s*<td><h4>Lecturer Name:.*$',
            '',
            html,
            flags=re.IGNORECASE | re.DOTALL
        )

        return html.strip()

    # =====================================================
    # LAP
    # =====================================================

    def get_lap(self):

        result = {

            "lapCode": None,

            "qualificationCode":
                self.get_qualification_code(),

            "clusterName":
                self.get_cluster_name(),

            "semester": None,

            "year": None,

            "lecturerName": None,

            "learningResources":
                self.get_learning_resources()
        }

        full_text = "\n".join(
            " ".join(row)
            for table in self.tables
            for row in table["rows"]
        )

        match = re.search(
            r"(20\d{2})\s+Semester\s+(\d)",
            full_text,
            re.IGNORECASE
        )

        if match:

            result["year"] = int(
                match.group(1)
            )

            result["semester"] = int(
                match.group(2)
            )

            return result

        year_match = re.search(
            r"Year\s*:?\s*(20\d{2})",
            full_text,
            re.IGNORECASE
        )

        semester_match = re.search(
            r"Semester\s*:?\s*(\d)",
            full_text,
            re.IGNORECASE
        )

        if year_match:
            result["year"] = int(
                year_match.group(1)
            )

        if semester_match:
            result["semester"] = int(
                semester_match.group(1)
            )

        return result

    # =====================================================
    # LECTURERS
    # =====================================================

    def get_lecturers(self):

        lecturers = []

        table = self.find_table_containing(
            "lecturer name"
        )

        if table is None:
            return lecturers

        rows = table["rows"]

        lecturer_header_found = False

        for row in rows:

            if len(row) < 5:
                continue

            if "lecturer name" in row[0].lower():

                lecturer_header_found = True
                continue

            if not lecturer_header_found:
                continue

            lecturer_name = row[0].strip()

            if not lecturer_name:
                continue

            room = row[4].strip()

            room = (
                room
                .replace("Joondalup", "")
                .replace("Perth", "")
                .replace("Midland", "")
                .strip()
            )

            lecturers.append({

                "lecturerName":
                    lecturer_name,

                "email":
                    row[2].strip(),

                "phone":
                    row[1].strip(),

                "room":
                    room,

                "contact":
                    row[3].strip()
            })

        return lecturers

    # =====================================================
    # ASSESSMENTS
    # =====================================================

    def get_assessment_type(self, title):

        value = title.lower()

        if "portfolio" in value:
            return "Portfolio"

        if "project" in value:
            return "Project"

        if "knowledge" in value:
            return "Knowledge Assessment"

        if "observation" in value:
            return "Observation Checklist"

        if "third party" in value:
            return "Third Party Report"

        if "workplace" in value:
            return "Workplace Assessment"

        if "rpl" in value:
            return "Recognition of Prior Learning"

        return "Other"

    def get_assessments(self):

        assessments = []

        for table in self.tables:

            for row in table["rows"]:

                if len(row) < 3:
                    continue

                if not row[0].startswith(
                    "Assessment"
                ):
                    continue

                match = re.search(
                    r"(\d+)",
                    row[0]
                )

                if not match:
                    continue

                title_text = row[1].strip()

                title = title_text
                description = ""

                if "–" in title_text:

                    title, description = (
                        title_text.split(
                            "–",
                            1
                        )
                    )

                elif "-" in title_text:

                    title, description = (
                        title_text.split(
                            "-",
                            1
                        )
                    )

                due_week = None

                week_match = re.search(
                    r"Week\s+(\d+)",
                    row[2],
                    re.IGNORECASE
                )

                if week_match:

                    due_week = int(
                        week_match.group(1)
                    )

                assessments.append({

                    "assessmentNumber":
                        int(match.group(1)),

                    "assessmentTitle":
                        title.strip(),

                    "assessmentDescription":
                        description.strip(),

                    "assessmentType":
                        self.get_assessment_type(
                            title_text
                        ),

                    "dueWeek":
                        due_week
                })

        return assessments

    # =====================================================
    # SESSION TABLE
    # =====================================================

    def is_session_table(self, table):

        text = "\n".join(
            " ".join(row)
            for row in table["rows"]
        ).lower()

        return (
            "session" in text
            and "topic" in text
            and "learning resources" in text
        )

    # =====================================================
    # SESSIONS
    # =====================================================

    def get_sessions(self):

        sessions = []

        for table in self.tables:

            if not self.is_session_table(table):
                continue

            for row in table["rows"]:

                try:
                    session_number = int(
                        str(row[0]).strip()
                    )
                except:
                    continue

                session_id = (
                    f"{self.lap['lapCode']}|"
                    f"{session_number}"
                )

                # UoC / Element mappings

                if len(row) > 2:

                    element_text = row[2]

                    references = re.findall(
                        r"[A-Z]{3,}\d{3,}(?:\.\d+)?",
                        element_text
                    )

                    for ref in references:

                        self.session_element_mappings.append({

                            "sessionId":
                                session_id,

                            "elementReference":
                                ref
                        })

                try:
                    face_to_face = float(
                        row[1] or 0
                    )
                except:
                    face_to_face = 0

                try:
                    out_of_class = float(
                        row[6] or 0
                    )
                except:
                    out_of_class = 0

                topic = (
                    row[3].strip()
                    if len(row) > 3
                    else ""
                )

                topic = re.sub(
                    r"^\s*Session\s+\d+\s*",
                    "",
                    topic,
                    flags=re.IGNORECASE
                ).strip()

                learning_resources = (
                    row[4].strip()
                    if len(row) > 4
                    else ""
                )

                out_of_class_activity = (
                    row[5].strip()
                    if len(row) > 5
                    else ""
                )

                assessment_hours = 0

                if (
                    "submission" in topic.lower()
                    or "assessment" in topic.lower()
                ):

                    assessment_hours = (
                        face_to_face
                    )

                    face_to_face = 0

                sessions.append({

                    "sessionId":
                        session_id,

                    "sessionNumber":
                        session_number,

                    "topic":
                        topic,

                    "faceToFaceHours":
                        face_to_face,

                    "workshopHours":
                        0,

                    "workPlacementHours":
                        0,

                    "outOfClassStudyHours":
                        out_of_class,

                    "tutorialSupportHours":
                        0,

                    "assessmentHours":
                        assessment_hours,

                    "learningResources":
                        learning_resources,

                    "outOfClassActivity":
                        out_of_class_activity
                })

        return sessions

    # =====================================================
    # PARSE
    # =====================================================

    def parse(self):

        self.lap = self.get_lap()

        qualification_code = (
            self.lap.get(
                "qualificationCode"
            )
        )

        cluster_name = (
            self.lap.get(
                "clusterName"
            )
        )

        cluster_key = (
            self.make_cluster_key(
                cluster_name
            )
        )

        if (
            qualification_code
            and cluster_key
        ):
            self.lap["lapCode"] = (
                f"{qualification_code}-"
                f"{cluster_key}"
            )

        self.lecturers = (
            self.get_lecturers()
        )

        if self.lecturers:

            self.lap["lecturerName"] = (
                self.lecturers[0][
                    "lecturerName"
                ]
            )

        self.assessments = (
            self.get_assessments()
        )

        self.sessions = (
            self.get_sessions()
        )

        return self

    # =====================================================
    # OUTPUT
    # =====================================================

    def to_lists(self):

        return {

            "lap":
                [self.lap],

            "sessions":
                self.sessions,

            "assessments":
                self.assessments,

            "lecturers":
                self.lecturers,

            "sessionelementmappings":
                self.session_element_mappings
        }