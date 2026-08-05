from datetime import datetime
from importlib.resources import path
from unittest import result

from docx import Document
from io import BytesIO
from docx.oxml.ns import qn
import re
import xml
import html
from lxml import etree


class WordHtmlConverter:

    NS = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "w14": "http://schemas.microsoft.com/office/word/2010/wordml",
    }

    def word_xml_to_html(self, xml):

        if xml is None:
            return ""

        #
        # Handle python-docx objects
        #
        if hasattr(xml, "xml"):
            xml = xml.xml

        if isinstance(xml, str):
            root = etree.fromstring(
                xml.encode("utf-8")
            )
        else:
            root = xml

        return self._process_element(root)

    def _process_element(self, element):

        tag = etree.QName(element).localname

        if tag == "tbl":
            return self._table_to_html(element)

        if tag == "tc":
            return self._cell_to_html(element)

        if tag == "p":
            return self._paragraph_to_html(element)

        parts = []

        for child in element:
            parts.append(
                self._process_element(child)
            )

        return "".join(parts)

    def _table_to_html(self, table):

        rows_html = []

        rows_html.append(
            '<table border="1">'
        )

        rows = table.xpath(
            "./w:tr",
            namespaces=self.NS
        )

        for row in rows:

            rows_html.append("<tr>")

            cells = row.xpath(
                "./w:tc",
                namespaces=self.NS
            )

            for cell in cells:

                rows_html.append(
                    (
                        "<td>"
                        f"{self._cell_to_html(cell)}"
                        "</td>"
                    )
                )

            rows_html.append("</tr>")

        rows_html.append("</table>")

        return "\n".join(rows_html)

    def _cell_to_html(self, cell):

        parts = []

        in_list = False

        paragraphs = cell.xpath(
            "./w:p",
            namespaces=self.NS
        )

        for paragraph in paragraphs:

            is_list = bool(
                paragraph.xpath(
                    "./w:pPr/w:numPr",
                    namespaces=self.NS
                )
            )

            html = self._paragraph_to_html(
                paragraph
            )

            if not html:
                continue

            if is_list:

                if not in_list:
                    parts.append("<ul>")
                    in_list = True

                parts.append(html)

            else:

                if in_list:
                    parts.append("</ul>")
                    in_list = False

                parts.append(html)

        if in_list:
            parts.append("</ul>")

        return "\n".join(parts)

    def _paragraph_to_html(self, paragraph):

        is_list = bool(
            paragraph.xpath(
                "./w:pPr/w:numPr",
                namespaces=self.NS
            )
        )

        runs = paragraph.xpath(
            "./w:r",
            namespaces=self.NS
        )

        #
        # Build formatting groups
        #
        groups = []

        for run in runs:

            text = "".join(
                run.xpath(
                    ".//w:t/text()",
                    namespaces=self.NS
                )
            )

            if not text:
                continue

            fmt = (
                bool(run.xpath(
                    "./w:rPr/w:b",
                    namespaces=self.NS
                )),
                bool(run.xpath(
                    "./w:rPr/w:i",
                    namespaces=self.NS
                )),
                bool(run.xpath(
                    "./w:rPr/w:u",
                    namespaces=self.NS
                ))
            )

            #
            # Merge adjacent runs with same formatting
            #
            if groups and groups[-1]["fmt"] == fmt:

                groups[-1]["text"] += text

            else:

                groups.append({
                    "fmt": fmt,
                    "text": text
                })

        if not groups:
            return ""

        #
        # Is entire paragraph bold?
        #
        all_bold = all(
            g["fmt"][0]
            for g in groups
            if g["text"].strip()
        )

        #
        # Heading detection
        #
        plain_text = "".join(
            g["text"]
            for g in groups
        ).strip()

        if (
            all_bold
            and plain_text
        ):
            return f"<h4>{html.escape(plain_text)}</h4>"

        #
        # Normal paragraph
        #
        parts = []

        for group in groups:

            text = html.escape(
                group["text"]
            )

            bold, italic, underline = (
                group["fmt"]
            )

            if underline:
                text = f"<u>{text}</u>"

            if italic:
                text = f"<em>{text}</em>"

            if bold:
                text = f"<strong>{text}</strong>"

            parts.append(text)

        content = "".join(parts).strip()

        if not content:
            return ""

        if is_list:
            return f"<li>{content}</li>"

        return f"<p>{content}</p>"


    def _run_to_html(self, run):

        #
        # Checkbox
        #
        checked = run.xpath(
            ".//w14:checked[@w14:val='1']",
            namespaces=self.NS
        )

        if checked:
            return "☑"

        text = "".join(
            run.xpath(
                ".//w:t/text()",
                namespaces=self.NS
            )
        )

        if not text:

            br = run.xpath(
                ".//w:br",
                namespaces=self.NS
            )

            if br:
                return "<br>"

            return ""

        text = html.escape(text)

        #
        # Formatting
        #
        if run.xpath(
            "./w:rPr/w:b",
            namespaces=self.NS
        ):
            text = (
                f"<strong>{text}</strong>"
            )

        if run.xpath(
            "./w:rPr/w:i",
            namespaces=self.NS
        ):
            text = (
                f"<em>{text}</em>"
            )

        if run.xpath(
            "./w:rPr/w:u",
            namespaces=self.NS
        ):
            text = (
                f"<u>{text}</u>"
            )

        return text

    def _hyperlink_to_html(self, hyperlink):

        text = ""

        for run in hyperlink.xpath(
            ".//w:r",
            namespaces=self.NS
        ):
            text += self._run_to_html(
                run
            )

        #
        # Relationship lookup could
        # be added later.
        #
        return text

class TASDoc:

    def __init__(self, source):

        self.source = source        
        if isinstance(source, bytes):
            self.source = BytesIO(source)        
        self.doc = Document(self.source)

        self.tables = self._load_tables()
        self.template = self._detect_template()

        self.converter = WordHtmlConverter()

        # Parsed data
        self.qualification = None
        self.delivery = None
        self.delivery_content = None
        self.units = None

        # Relationship data
        self.clusters = {}
        self.cluster_units = {}
        self.qualification_clusters = {}
        self.qualification_units = {}


    # =====================================================
    # CORE
    # =====================================================

    @staticmethod
    def clean(text):

        if text is None:
            return ""

        return " ".join(str(text).split())

    def _load_tables(self):

        tables = []

        for table_number, table in enumerate(self.doc.tables):

            rows = []

            for row in table.rows:

                rows.append([
                    self.clean(cell.text)
                    for cell in row.cells
                ])

            tables.append({
                "table_number": table_number,
                "table": table,
                "rows": rows
            })

        return tables

    def _detect_template(self):

        for table in self.tables:

            rows = table["rows"]

            for row in rows:

                text = " ".join(row)

                if (
                    "This TAS version is current as at:"
                    in text
                ):
                    return "new"

                if (
                    "Training package code and title"
                    in text
                    and len(self.tables) > 28
                ):
                    return "new"

        return "old"

    # =====================================================
    # QUALIFICATION
    # =====================================================
    def get_qualification(self):

        for table in self.tables:
            
            table_xml = str(table["table"]._tbl.xml)

            rows = table["rows"]

            for i, row in enumerate(rows):

                text = " ".join(row).lower()

                if (
                    "qualification national code and title"
                    in text
                ):

                    if i < 1 or i + 1 >= len(rows):
                        continue

                    training_package_row = rows[i - 1]
                    qualification_row = rows[i + 1]

                    # Training package

                    package_value = (
                        training_package_row[0]
                        .strip()
                    )

                    if " - " in package_value:

                        package_code, package_title = (
                            package_value.split(
                                " - ",
                                1
                            )
                        )

                    else:

                        package_code = None
                        package_title = package_value

                    # Qualification

                    qualification_value = (
                        qualification_row[0]
                        .strip()
                    )

                    qualification_parts = (
                        qualification_value.split(
                            " ",
                            1
                        )
                    )

                    if len(qualification_parts) == 2:

                        qualification_code = (
                            qualification_parts[0]
                        )

                        qualification_title = (
                            qualification_parts[1]
                        )                        

                    else:

                        qualification_code = None
                        qualification_title = (
                            qualification_value
                        )

                    qualification_title = (
                        qualification_title
                        .lstrip("-")
                        .strip()
                    )

                    return {

                        "training_package": {

                            "code":
                                package_code,

                            "title":
                                package_title,

                            "release_number":
                                training_package_row[1]
                                if len(training_package_row) > 1
                                else None,

                            "release_date":
                                training_package_row[2]
                                if len(training_package_row) > 2
                                else None
                        },

                        "qualification": {

                            "national_code":
                                qualification_code,

                            "title":
                                qualification_title,

                            "release_number":
                                qualification_row[1]
                                if len(qualification_row) > 1
                                else None,

                            "release_date":
                                qualification_row[2]
                                if len(qualification_row) > 2
                                else None,

                            "state_code":
                                qualification_row[3]
                                if len(qualification_row) > 3
                                else None,

                            "on_scope":
                                ( 'w14:checked w14:val="1"' in table_xml)
                        }
                    }

        return {}

    # =====================================================
    # DELIVERY DETAILS
    # =====================================================

    def get_qualification_classification(self):

        for table in self.tables:

            table_text = "\n".join(
                " ".join(row)
                for row in table["rows"]
            ).lower()

            if (
                "qualification classification"
                not in table_text
            ):
                continue

            docx_table = table["table"]

            mapping = {
                1: "A",
                2: "B",
                3: "C"
            }

            for row in docx_table.rows:

                for col, classification in mapping.items():

                    if col >= len(row.cells):
                        continue

                    if (
                        'w14:checked w14:val="1"'
                        in row.cells[col]._tc.xml
                    ):
                        return classification

        return None

    def get_enrolment_type(self):

        if self.template == "new":
            return self._get_enrolment_type_new()

        return self._get_enrolment_type_old()
    
    def _get_enrolment_type_new(self):

        for table in self.tables:

            rows = table["rows"]

            if not rows:
                continue

            table_text = " ".join(
                " ".join(row)
                for row in rows
            ).lower()

            if (
                "full time" in table_text
                and "part time" in table_text
            ):

                docx_table = table["table"]

                checkbox_row = docx_table.rows[0]

                #
                # Full Time
                #

                if (
                    'w14:checked w14:val="1"'
                    in checkbox_row.cells[0]._tc.xml
                ):
                    return "Full Time"

                #
                # Part Time
                #

                if (
                    'w14:checked w14:val="1"'
                    in checkbox_row.cells[1]._tc.xml
                ):
                    return "Part Time"

                #
                # Other
                #

                if (
                    'w14:checked w14:val="1"'
                    in checkbox_row.cells[2]._tc.xml
                ):

                    return self.get_other_enrolment_value(
                        checkbox_row.cells[2]
                    )

                #
                # Third Party
                #

                if (
                    'w14:checked w14:val="1"'
                    in checkbox_row.cells[3]._tc.xml
                ):
                    return "Third Party"

        return None

    def _get_enrolment_type_old(self):

        for table in self.tables:

            text = "\n".join(
                " ".join(row)
                for row in table["rows"]
            ).lower()

            if (
                "full time" not in text
                or "part time" not in text
            ):
                continue

            xml = table["table"]._tbl.xml

            checked_values = re.findall(
                r'w14:checked w14:val="([01])"',
                xml
            )

            if len(checked_values) < 4:
                continue

            mapping = [
                "Full Time",
                "Part Time",
                "Other",
                "Third Party"
            ]

            for i, value in enumerate(
                checked_values[:4]
            ):

                if value == "1":

                    if mapping[i] == "Other":

                        return self.get_other_enrolment_value(
                            table["table"].cell(0, 0)
                        )

                    return mapping[i]

        return None

    def get_other_enrolment_value(self, cell):

        xml = cell._tc.xml

        #
        # Find dropdown content controls
        #
        dropdowns = re.findall(
            r'<w:sdt>.*?<w:dropDownList>.*?</w:dropDownList>.*?<w:t>(.*?)</w:t>',
            xml,
            flags=re.DOTALL
        )

        for value in dropdowns:

            value = value.strip()

            if not value:
                continue

            if value.lower() == "click here to select other type":
                continue

            return value

        return "Other"

    def get_delivery(self):

        result = {
            "title": None,
            "qualification_state_code": None,
            "campus": None,
            "campus_code": None,
            "enrolment_type": self.get_enrolment_type(),
            "enrolment_code": None,
            "duration": None,
            "start_date": None,
            "finish_date": None,
            "semester": None,
            "year": None,
            "stages": None,
            "template_version": None,
            "version_number": None,
            "approval_status": None
        }

        for table in self.tables:

            for row in table["rows"]:

                for i in range(len(row) - 1):

                    label = (
                        row[i]
                        .strip()
                        .lower()
                        .rstrip(":")
                    )

                    value = row[i + 1].strip()

                    #
                    # Duration / Dates
                    #
                    if label == "duration":

                        result["duration"] = value

                        xml = table["table"]._tbl.xml

                        dates = re.findall(
                            r'w:fullDate="([^"]+)"',
                            xml
                        )

                        if len(dates) >= 2:

                            from datetime import datetime

                            result["start_date"] = (
                                datetime.fromisoformat(
                                    dates[0].replace(
                                        "Z",
                                        "+00:00"
                                    )
                                ).date()
                            )

                            result["finish_date"] = (
                                datetime.fromisoformat(
                                    dates[1].replace(
                                        "Z",
                                        "+00:00"
                                    )
                                ).date()
                            )

                    #
                    # Campus
                    #
                    elif (
                        "campus" in label
                        or "delivery location" in label
                        or "delivery site" in label
                    ):

                        result["campus"] = value
                    
                    #
                    # Number of Stages
                    #
                    elif label in [
                        "number of stages",
                        "stages",
                        "no. of stages"
                    ]:

                        try:
                            result["stages"] = int(value)
                        except ValueError:
                            pass

        #
        # Derived Values
        #

        start_date = result["start_date"]

        if start_date:

            result["year"] = start_date.year

            result["semester"] = (
                "S1"
                if start_date.month <= 6
                else "S2"
            )

        #
        # Campus Code
        #

        campus = (
            result["campus"] or ""
        ).lower()

        if "joondalup" in campus:
            result["campus_code"] = "J"

        elif "perth" in campus:
            result["campus_code"] = "P"

        elif "midland" in campus:
            result["campus_code"] = "M"

        elif "clarkson" in campus:
            result["campus_code"] = "C"

        else:
            result["campus_code"] = "UNK"

        #
        # Enrolment Code
        #

        enrolment = (
            result["enrolment_type"] or ""
        ).lower()

        if "full" in enrolment:
            result["enrolment_code"] = "FT"

        elif "part" in enrolment:
            result["enrolment_code"] = "PT"

        elif "apprent" in enrolment:
            result["enrolment_code"] = "APP"

        elif "trainee" in enrolment:
            result["enrolment_code"] = "TRN"

        elif "third" in enrolment: 
            result["enrolment_code"] = "TP"

        elif "pre-apprent" in enrolment:
            result["enrolment_code"] = "PREAPP"

        elif "vet in schools" in enrolment:
            result["enrolment_code"] = "VETIS"

        elif "fee" in enrolment:
            result["enrolment_code"] = "FFS"
            
        else:
            result["enrolment_code"] = "UNK"

        #
        # Template Version
        #

        result["template_version"] = (
            "11.0"
            if self.template == "new"
            else "10.0"
        )

        #
        # Delivery Title
        #

        if self.qualification:

            qual = self.qualification.get(
                "qualification",
                {}
            )

            national_code = qual.get(
                "national_code",
                "UNKNOWN"
            )

            state_code = qual.get(
                "state_code",
                "UNKNOWN"
            )

            result[
                "qualification_state_code"
            ] = state_code

            if (
                result["year"]
                and result["semester"]
            ):

                result["title"] = (
                    f"{national_code}-"
                    f"{state_code}-"
                    f"{result['campus_code']}-"
                    f"{result['enrolment_code']}-"
                    f"{result['year']}-"
                    f"{result['semester']}"
                )

        return result
    
    def _get_table_by_number(self, table_number):

        for table in self.tables:

            if (
                table["table_number"]
                == table_number
            ):
                return table

        return None

    def _get_delivery_content_table_map(self):

        return {
            "old": {

                "program_overview": 5,
                "learner_cohort": 6,

                "enrolment": 7,
                "duration": 8,
                "campus": 9,

                "delivery_rationale": 11,
                "amount_of_training": 12,

                "learning_resources": 23,
                "facilities_equipment": 24,

                "learner_support": 26,
                "pathways": 27,
            },

            "new": {

                "program_overview": 4,
                "learner_cohort": 5,

                "enrolment": 6,
                "duration": 7,
                "campus": 8,

                "delivery_rationale": 10,
                "amount_of_training": 11,

                "learning_resources": 24,
                "facilities_equipment": 25,

                "learner_support": 26,
                "pathways": 28,

                "continuous_improvement": 29,
                "review_information": 30,
            }
        }

    def get_delivery_content(self):

        result = {

            #
            # HTML
            #
            "program_overview": None,
            "industry_engagement": None,
            "learner_cohort": None,
            "delivery_rationale": None,
            "amount_of_training": None,
            "learning_resources": None,
            "facilities_equipment": None,
            "learner_support": None,
            "pathways": None,
            "continuous_improvement": None,

            #
            # Plain text
            #
            "program_overview_text": None,
            "industry_engagement_text": None,
            "learner_cohort_text": None,
            "delivery_rationale_text": None,
            "amount_of_training_text": None,
            "learning_resources_text": None,
            "facilities_equipment_text": None,
            "learner_support_text": None,
            "pathways_text": None,
            "continuous_improvement_text": None,

            #
            # Metadata
            #
            "qualification_classification":
                self.get_qualification_classification(),

            "industry_meeting_date": None,
            "review_date": None,
            "current_as_at_date": None
        }

        mapping = self._get_delivery_content_table_map()[
            self.template
        ]

        #
        # Generic extraction
        #
        for section, table_number in mapping.items():

            table = self._get_table_by_number(
                table_number
            )

            if table is None:
                continue

            text = "\n".join(
                " ".join(row)
                for row in table["rows"]
            )

            html = self.converter.word_xml_to_html(
                table["table"]._tbl
            )

            #
            # Store HTML
            #
            if section in result:
                result[section] = html

            #
            # Store text
            #
            text_key = f"{section}_text"

            if text_key in result:
                result[text_key] = text

        #
        # NEW TEMPLATE SPECIAL HANDLING
        #
        if self.template == "new":

            #
            # Program Overview / Industry Engagement
            #
            overview_table = self._get_table_by_number(
                mapping["program_overview"]
            )

            if overview_table:

                text = "\n".join(
                    " ".join(row)
                    for row in overview_table["rows"]
                )

                html = self.converter.word_xml_to_html(
                    overview_table["table"]._tbl
                )

                #
                # Newer templates explicitly split
                #
                if re.search(
                    r"industry engagement and feedback:",
                    text,
                    flags=re.IGNORECASE
                ):

                    text_parts = re.split(
                        r"industry engagement and feedback:",
                        text,
                        flags=re.IGNORECASE
                    )

                    html_parts = re.split(
                        r"industry engagement and feedback:",
                        html,
                        flags=re.IGNORECASE
                    )

                    result[
                        "program_overview_text"
                    ] = text_parts[0].strip()

                    result[
                        "program_overview"
                    ] = html_parts[0].strip()

                    if len(text_parts) > 1:

                        result[
                            "industry_engagement_text"
                        ] = text_parts[1].strip()

                    if len(html_parts) > 1:

                        result[
                            "industry_engagement"
                        ] = html_parts[1].strip()

                #
                # Some TAS documents don't contain the
                # Industry Engagement heading at all.
                #
                else:

                    result[
                        "program_overview_text"
                    ] = text

                    result[
                        "program_overview"
                    ] = html

                    #
                    # Treat the whole section as
                    # industry engagement too.
                    #
                    result[
                        "industry_engagement_text"
                    ] = text

                    result[
                        "industry_engagement"
                    ] = html

                #
                # Industry meeting date
                #
                match = re.search(
                    r"industry advisory committee met on\s+(\d{1,2}/\d{1,2}/\d{4})",
                    text,
                    flags=re.IGNORECASE
                )

                if match:

                    result[
                        "industry_meeting_date"
                    ] = match.group(1)

                #
                # Merge learner support
                #
                learner_support = result.get(
                    "learner_support_text"
                )

                extra_table = self._get_table_by_number(27)

                if extra_table:

                    extra_text = "\n".join(
                        " ".join(row)
                        for row in extra_table["rows"]
                    )

                    extra_html = self.converter.word_xml_to_html(
                        extra_table["table"]._tbl
                    )

                    extra_text = extra_text.replace(
                        (
                            "Amend this list as appropriate "
                            "to the learner cohort remove "
                            "this sentence and amend the "
                            "font colour to black."
                        ),
                        ""
                    ).strip()

                    if learner_support:

                        result[
                            "learner_support_text"
                        ] = (
                            learner_support
                            + "\n\n"
                            + extra_text
                        )

                    if result["learner_support"]:

                        result[
                            "learner_support"
                        ] += (
                            "<br><br>"
                            + extra_html
                        )

                #
                # Review information
                #
                review_table = self._get_table_by_number(
                    30
                )

                if review_table:
                    
                    from datetime import datetime

                    for row in review_table["rows"]:

                        row_text = " ".join(
                            str(cell)
                            for cell in row
                        )

                        if (
                            "reviewed"
                            in row_text.lower()
                        ):

                            match = re.search(
                                r"(\d{1,2}/\d{1,2}/\d{4})",
                                row_text
                            )

                            if match:

                                result[
                                    "review_date"
                                ] = (   datetime.strptime(
                                        match.group(1),
                                        "%d/%m/%Y"
                                        ).date()
                                    )
                        elif (
                            "current as at"
                            in row_text.lower()
                        ):

                            match = re.search(
                                r"(\d{1,2}/\d{1,2}/\d{4})",
                                row_text
                            )

                            if match:

                                result[
                                    "current_as_at_date"
                                ] = (   datetime.strptime(
                                        match.group(1),
                                        "%d/%m/%Y"
                                        ).date()
                                    )

        return result

    # =====================================================
    # DELIVERY SCHEDULE
    # =====================================================

    def _is_delivery_schedule(self, rows):

        if len(rows) < 2:
            return False

        header_text = " ".join(
            " ".join(row)
            for row in rows[:2]
        ).lower()

        return (
            "cluster" in header_text
            and "unit of competency" in header_text
            and "nominal hours" in header_text
        )    

    def parse_core_elective(self, value):

        value = value.strip().upper()

        if value == "C":
            return "C", None

        if "IMPORT" in value:
            return "E", "IMPORT"

        match = re.search(
            r"E\s*\(([A-Z])\)",
            value
        )

        if match:
            return "E", match.group(1)

        return "E", None

    def get_units(self):

        units = []

        stage = 1

        for table in self.tables:

            rows = table["rows"]
            docx_table = table["table"]            
            cluster_map = self.build_cluster_map(
                    docx_table
                )


            if not self._is_delivery_schedule(rows):
                continue

            if units:
                stage += 1

            for row_index, row in enumerate(rows):

                if len(row) < 4:
                    continue

                code = row[1].strip()

                if not re.fullmatch(
                    r"[A-Z]{3,}\d{3,}",
                    code
                ):
                    continue

                # column mappings differ slightly
                if self.template == "new":

                    delivery_mode = (
                        row[13]
                        if len(row) > 13
                        else ""
                    )

                    core_elective_raw = (
                        row[14]
                        if len(row) > 14
                        else ""
                    )
                    
                    core_elective, elective_group = (
                        self.parse_core_elective(
                            core_elective_raw
                        )
)


                else:

                    delivery_mode = (
                        row[7]
                        if len(row) > 7
                        else ""
                    )

                    core_elective_raw = (
                        row[8]
                        if len(row) > 8
                        else ""
                    )

                    core_elective, elective_group = (
                        self.parse_core_elective(
                            core_elective_raw
                        )
                    )
                
                units.append({
                    "stage": stage,
                    "cluster":(
                                " ".join(cluster_map[row_index].split())
                                if cluster_map.get(row_index)
                                else None
                            ),
                    "national_code": code,
                    "state_code": row[2].strip(),
                    "unit_title": row[3].strip(),
                    "nominal_hours": (
                        row[4]
                        if len(row) > 4
                        else ""
                    ),
                    "delivery_mode": delivery_mode,
                    "core_elective": core_elective,
                    "elective_group": elective_group
                })                

        return units
        
    def get_cell_borders(self, cell):

        tc_pr = cell._tc.tcPr

        if tc_pr is None:
            return {}

        borders = tc_pr.find(qn("w:tcBorders"))

        if borders is None:
            return {}

        result = {}

        for side in ["top", "bottom", "left", "right"]:

            border = borders.find(qn(f"w:{side}"))

            if border is None:

                result[side] = None

            else:

                result[side] = border.get(
                    qn("w:val")
                )

        return result
        
    def get_cluster_state(self, cell):

        borders = self.get_cell_borders(cell)

        top = borders.get("top")
        bottom = borders.get("bottom")

        if top == "single" and bottom == "nil":
            return "start"

        if top == "nil" and bottom == "nil":
            return "middle"

        if top == "nil" and bottom == "single":
            return "end"

        if top == "single" and bottom == "single":
            return "standalone"

        return "unknown"
    
    def build_cluster_map(self, docx_table):

        cluster_map = {}

        current_cluster_parts = []

        current_rows = []

        for row_idx, row in enumerate(docx_table.rows):

            if len(row.cells) < 2:
                continue

            code = row.cells[1].text.strip()

            #
            # Only process actual unit rows
            #
            if not re.fullmatch(
                r"[A-Z]{3,}\d{3,}",
                code
            ):
                continue

            text = row.cells[0].text.strip()

            state = self.get_cluster_state(
                row.cells[0]
            )

            #
            # START OF CLUSTER
            #
            if state == "start":

                #
                # Flush previous unfinished cluster
                #
                if current_rows:

                    cluster_name = " ".join(
                        current_cluster_parts
                    ).strip()

                    for r in current_rows:
                        cluster_map[r] = cluster_name

                current_cluster_parts = []

                current_rows = [row_idx]

                if text:
                    current_cluster_parts.append(text)

            #
            # MIDDLE OF CLUSTER
            #
            elif state == "middle":

                current_rows.append(row_idx)

                if text:
                    current_cluster_parts.append(text)

            #
            # END OF CLUSTER
            #
            elif state == "end":

                #
                # End of multi-row cluster
                #
                if current_rows:

                    current_rows.append(row_idx)

                    if text:
                        current_cluster_parts.append(text)

                    cluster_name = " ".join(
                        current_cluster_parts
                    ).strip()

                    for r in current_rows:
                        cluster_map[r] = cluster_name

                    current_cluster_parts = []

                    current_rows = []

                #
                # Standalone row
                #
                else:

                    cluster_map[row_idx] = (
                        text if text else None
                    )

            #
            # STANDALONE CLUSTER
            #
            elif state == "standalone":

                cluster_map[row_idx] = (
                    text if text else None
                )

            #
            # UNKNOWN STATE
            #
            else:

                cluster_map[row_idx] = (
                    text if text else None
                )

        #
        # Flush unfinished cluster at end of table
        #
        if current_rows:

            cluster_name = " ".join(
                current_cluster_parts
            ).strip()

            for r in current_rows:
                cluster_map[r] = cluster_name

        return cluster_map      

    def make_cluster_key(self, cluster_name):

        cluster_name = html.unescape(cluster_name)

        cluster_name = re.sub(
            r'([a-zA-Z])(\d+)',
            r'\1-\2',
            cluster_name
        )

        cluster_name = cluster_name.lower()

        cluster_name = cluster_name.replace("&", "and")

        cluster_name = re.sub(
            r'[^a-z0-9]+',
            '-',
            cluster_name
        )

        cluster_name = re.sub(
            r'-+',
            '-',
            cluster_name
        )

        return cluster_name.strip("-")
    
    def make_folder_name(self, cluster_name):

        joining_words = {
            "and",
            "or",
            "of",
            "the",
            "for",
            "to",
            "in",
            "on",
            "at",
            "by",
            "with"
        }

        cluster_name = html.unescape(cluster_name)

        # Separate letters and numbers
        cluster_name = re.sub(
            r'([a-zA-Z])(\d+)',
            r'\1 \2',
            cluster_name
        )

        # Replace special chars with separator
        cluster_name = re.sub(
            r'[^A-Za-z0-9 ]+',
            ' - ',
            cluster_name
        )

        # Collapse multiples
        cluster_name = re.sub(
            r'(?:\s*-\s*)+',
            ' - ',
            cluster_name
        )

        cluster_name = ' '.join(cluster_name.split())

        words = []

        for i, word in enumerate(cluster_name.split()):

            if word == "-":
                words.append(word)
                continue

            lower = word.lower()
            
            # Preserve acronyms adn oddly cased words
            if any(c.isupper() for c in word[1:]):
                words.append(word)
                continue

            if (
                i > 0
                and lower in joining_words
            ):
                words.append(lower)
            else:
                words.append(lower.capitalize())

        return " ".join(words)
    
    def make_cluster_unit_id(self, cluster_key, state_code):

        parts = re.findall(
            r'(?:^|[-_])([a-z])|(?:^|[-_])(\d+)',
            cluster_key
        )

        acronym = ''.join(
            letter or number
            for letter, number in parts
        ).upper()


        return f"{acronym}-{state_code}"

    
    def build_relationships(self):

        self.clusters = {}
        self.cluster_units = {}
        self.qualification_clusters = {}

        qualification_code = self.qualification["qualification"]["state_code"]

        for unit in self.units:

            #
            # Qualification -> Unit
            #
            self.qualification_units[
                (qualification_code, unit["state_code"])
            ] = {
                "qualificationUnitId": f"{qualification_code}-{unit['state_code']}",
                "qualificationCode": qualification_code,
                "unitCode": unit["state_code"],
                "coreElective": unit["core_elective"],
                "electiveGroup": unit["elective_group"]
            }

            cluster_name = unit["cluster"]

            if not cluster_name:
                continue

            cluster_key = self.make_cluster_key(cluster_name)
            cluster_folder = self.make_folder_name(cluster_name)

            self.clusters[cluster_key] = {
                "clusterKey": cluster_key,
                "clusterName": cluster_name,
                "clusterFolder": cluster_folder
            }

            self.qualification_clusters[
                (qualification_code, cluster_key)
            ] = {
                "qualificationCode": qualification_code,
                "clusterKey": cluster_key,
                "stage": unit["stage"]
            }

            self.cluster_units[
                (cluster_key, unit["state_code"])
            ] = {
                "clusterUnitId": self.make_cluster_unit_id
                    (cluster_key, unit["state_code"]),
                "clusterKey": cluster_key,
                "stateCode": unit["state_code"]
            }            
            

    # =====================================================
    # MASTER OUTPUT
    # =====================================================

    def parse(self):
        
        self.qualification = self.get_qualification()
        self.delivery = self.get_delivery()
        self.delivery_content = self.get_delivery_content()
        self.units = self.get_units()
        if (
            self.delivery.get("stages") is None
            and self.units
        ):

            stages = {
                unit["stage"]
                for unit in self.units
                if unit.get("stage")
            }

            if stages:
                self.delivery["stages"] = max(stages)                
        
        self.build_relationships()
        
        return self
    
    def to_lists(self):

        return {
            "template": self.template,
            "qualification": self.qualification,
            "delivery": self.delivery,
            "delivery_content": self.delivery_content,
            "units": self.units,
            "clusters": list(self.clusters.values()),
            "qualification_clusters": list(
                self.qualification_clusters.values()
            ),
            "cluster_units": list(
                self.cluster_units.values()
            ),
            "qualification_units": list(
                self.qualification_units.values()
            )         
        }